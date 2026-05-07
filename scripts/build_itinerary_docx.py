from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(r"C:\Users\aleja\Personal\Proyectos Web\barbersoft-react\ITINERARIO_PROYECTO_REACT_NODEJS.docx")


weeks = [
    {
        "semana": "1",
        "objetivo": "Base técnica, accesos y arquitectura inicial",
        "angel": [
            ("Frontend", "Crear proyecto React con Vite, routing, layout general y navegacion Admin/Empleado."),
            ("Backend", "Inicializar API Node.js, estructura de carpetas, variables de entorno y endpoint de salud."),
        ],
        "adrian": [
            ("Frontend", "Maquetar pantalla principal, login admin, login empleado y menus por rol."),
            ("Backend", "Definir modelo inicial de usuarios, roles y permisos para autenticación."),
        ],
        "miguel": [
            ("Frontend", "Definir sistema visual base: botones, inputs, tablas, estados y mensajes."),
            ("Backend", "Diseñar esquema relacional inicial: usuarios, empleados, clientes, servicios, productos y ventas."),
        ],
    },
    {
        "semana": "2",
        "objetivo": "Catálogos base de productos y servicios",
        "angel": [
            ("Frontend", "Conectar vistas de categorias con servicios de datos y manejo de carga/error."),
            ("Backend", "Crear rutas Node.js para categorías de productos y categorías de servicios."),
        ],
        "adrian": [
            ("Frontend", "Construir pantallas de catálogo de categorías de productos y categorías de servicios."),
            ("Backend", "Implementar validaciones de nombre único, estado activo/inactivo y respuestas de error."),
        ],
        "miguel": [
            ("Frontend", "Revisar consistencia de tablas, búsqueda, estados vacíos y formularios de catálogos."),
            ("Backend", "Crear seeds o datos iniciales para categorías y servicios de la barbería."),
        ],
    },
    {
        "semana": "3",
        "objetivo": "CRUD de productos y servicios",
        "angel": [
            ("Frontend", "Integrar formularios reutilizables para altas, ediciones y bajas lógicas."),
            ("Backend", "Implementar endpoints CRUD de productos con stock actual, stock mínimo, costo y precio."),
        ],
        "adrian": [
            ("Frontend", "Desarrollar CRUD de productos y CRUD de servicios con filtros y confirmaciones."),
            ("Backend", "Implementar endpoints CRUD de servicios con nombre, precio, descripción y categoría."),
        ],
        "miguel": [
            ("Frontend", "Probar flujos de productos/servicios, validaciones visuales y mensajes claros."),
            ("Backend", "Agregar validaciones de datos, control de estados activo/inactivo y pruebas de endpoints."),
        ],
    },
    {
        "semana": "4",
        "objetivo": "CRUD de clientes, empleados, perfil y permisos",
        "angel": [
            ("Frontend", "Crear hooks/servicios para consumir API de clientes, empleados y perfil."),
            ("Backend", "Implementar autenticación, middleware de roles y protección de rutas administrativas."),
        ],
        "adrian": [
            ("Frontend", "Construir CRUD de clientes, empleados y pantalla de perfil de usuario."),
            ("Backend", "Crear endpoints CRUD de clientes y empleados con altas, modificaciones y bajas."),
        ],
        "miguel": [
            ("Frontend", "Validar formularios de datos personales, permisos visibles y navegacion segun rol."),
            ("Backend", "Registrar bitácora de acciones críticas: alta, baja, edición, cancelación y login."),
        ],
    },
    {
        "semana": "5",
        "objetivo": "Venta de servicios e historial de visitas",
        "angel": [
            ("Frontend", "Implementar formulario de venta de servicio con cliente, empleado, servicio y metodo de pago."),
            ("Backend", "Crear endpoint transaccional para registrar servicio, venta, visita y comisión."),
        ],
        "adrian": [
            ("Frontend", "Agregar modal de alta rápida de cliente condicionado a la prestación del servicio."),
            ("Backend", "Validar reglas de negocio: cliente asociado, servicio activo y empleado autorizado."),
        ],
        "miguel": [
            ("Frontend", "Probar cálculos de subtotal, descuentos, total y mensajes de error del formulario."),
            ("Backend", "Guardar historial de visitas del cliente y preparar consulta para conteo de visitas."),
        ],
    },
    {
        "semana": "6",
        "objetivo": "Venta de productos e inventario",
        "angel": [
            ("Frontend", "Implementar formulario de venta de producto y venta adicional ligada a un servicio."),
            ("Backend", "Crear endpoint de venta de producto con actualización automática de stock."),
        ],
        "adrian": [
            ("Frontend", "Mostrar disponibilidad, stock mínimo, productos activos y bloqueo por falta de inventario."),
            ("Backend", "Validar existencia de producto, stock suficiente y cálculo de totales por línea."),
        ],
        "miguel": [
            ("Frontend", "Realizar pruebas de inventario, mensajes de stock bajo y resumen de productos vendidos."),
            ("Backend", "Agregar consultas de productos vendidos y auditoría de movimientos de inventario."),
        ],
    },
    {
        "semana": "7",
        "objetivo": "Consultas de ventas, rifa y comisiones",
        "angel": [
            ("Frontend", "Construir pantalla Mis Ventas con filtros por fecha, tipo y totales del empleado."),
            ("Backend", "Crear endpoints para consultar ventas del empleado autenticado y detalle de comisiones."),
        ],
        "adrian": [
            ("Frontend", "Agregar interfaz de registro de clientes en rifa y consulta de clientes elegibles."),
            ("Backend", "Implementar reglas de rifa basadas en historial de visitas y registro de participaciones."),
        ],
        "miguel": [
            ("Frontend", "Validar tablas de consulta, paginación, filtros y exportación sencilla de resultados."),
            ("Backend", "Implementar cálculo de comisiones por empleado y pruebas de exactitud financiera."),
        ],
    },
    {
        "semana": "8",
        "objetivo": "Gestión administrativa de ventas",
        "angel": [
            ("Frontend", "Integrar vista admin de ventas con detalle, edición controlada y estados de venta."),
            ("Backend", "Crear endpoints admin para listar, editar y cancelar ventas con permisos."),
        ],
        "adrian": [
            ("Frontend", "Construir diálogos de confirmación para cancelar/eliminar y formularios de corrección."),
            ("Backend", "Asegurar reglas transaccionales al editar/cancelar ventas y revertir stock si aplica."),
        ],
        "miguel": [
            ("Frontend", "Probar casos de venta de servicio, producto, producto adicional y cancelación."),
            ("Backend", "Registrar bitácora de cambios de venta y proteger historiales contra edición manual."),
        ],
    },
    {
        "semana": "9",
        "objetivo": "Corte de caja, reportes e indicadores",
        "angel": [
            ("Frontend", "Crear pantalla de corte de caja con totales por metodo de pago y efectivo declarado."),
            ("Backend", "Implementar endpoint de corte diario/turno con ingresos por servicios y productos."),
        ],
        "adrian": [
            ("Frontend", "Construir reportes de ingresos, ganancias, comisiones y filtros por rango de fechas."),
            ("Backend", "Crear consultas agregadas para reportes en menos de 2 segundos."),
        ],
        "miguel": [
            ("Frontend", "Diseñar tarjetas KPI, tablas resumidas y visualización clara de diferencias de caja."),
            ("Backend", "Validar exactitud de cálculos financieros, descuentos, comisiones y ganancias."),
        ],
    },
    {
        "semana": "10",
        "objetivo": "Integración, pruebas completas y entrega",
        "angel": [
            ("Frontend", "Integrar rutas finales, limpiar servicios de API y preparar flujo completo de demostración."),
            ("Backend", "Revisar configuración de Node.js, manejo centralizado de errores y seguridad básica."),
        ],
        "adrian": [
            ("Frontend", "Corregir pendientes de CRUD, responsive, navegación y consistencia de botones/menús."),
            ("Backend", "Completar pruebas de endpoints, datos de prueba y documentación de rutas."),
        ],
        "miguel": [
            ("Frontend", "Ejecutar QA final con evidencias de login, ventas, inventario, reportes y caja."),
            ("Backend", "Documentar despliegue, variables de entorno, respaldo de base de datos y cierre técnico."),
        ],
    },
]


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_in):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_in * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def clear_cell(cell):
    for paragraph in cell.paragraphs:
        paragraph._element.getparent().remove(paragraph._element)


def set_cell_text(cell, parts, header=False):
    clear_cell(cell)
    if isinstance(parts, str):
        p = cell.add_paragraph()
        run = p.add_run(parts)
        run.bold = header
        run.font.size = Pt(8.5 if header else 7.5)
        if header:
            run.font.color.rgb = RGBColor(255, 255, 255)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return

    for label, text in parts:
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.0
        label_run = p.add_run(f"[{label}] ")
        label_run.bold = True
        label_run.font.size = Pt(7.5)
        label_run.font.color.rgb = RGBColor(31, 78, 121)
        text_run = p.add_run(text)
        text_run.font.size = Pt(7.5)


def style_table(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [0.55, 1.8, 2.6, 2.6, 2.6]
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_width(cell, widths[idx])
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                for run in paragraph.runs:
                    run.font.name = "Aptos"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    for cell in table.rows[0].cells:
        shade_cell(cell, "1F4E79")
        tr_pr = table.rows[0]._tr.get_or_add_trPr()
        tbl_header = OxmlElement("w:tblHeader")
        tbl_header.set(qn("w:val"), "true")
        tr_pr.append(tbl_header)


def main():
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)

    title = doc.add_paragraph()
    title.text = "ITINERARIO DEL PROYECTO REACT + NODE.JS"
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(6)
    for run in title.runs:
        run.bold = True
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(31, 78, 121)

    intro = doc.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.CENTER
    intro_run = intro.add_run(
        "Planeación basada en los requerimientos de BarberSoft. "
        "Cada actividad indica si corresponde a Frontend (React) o Backend (Node.js)."
    )
    intro_run.font.size = Pt(9)
    intro_run.italic = True
    intro.paragraph_format.space_after = Pt(8)

    table = doc.add_table(rows=len(weeks) + 1, cols=5)
    table.style = "Table Grid"

    headers = ["Semana", "Objetivo", "Ángel", "Adrián", "Miguel"]
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, header=True)

    for row_idx, item in enumerate(weeks, start=1):
        cells = table.rows[row_idx].cells
        set_cell_text(cells[0], item["semana"])
        set_cell_text(cells[1], item["objetivo"])
        set_cell_text(cells[2], item["angel"])
        set_cell_text(cells[3], item["adrian"])
        set_cell_text(cells[4], item["miguel"])
        if row_idx % 2 == 0:
            for cell in cells:
                shade_cell(cell, "F3F6FA")

    style_table(table)
    doc.save(str(OUTPUT))
    print(OUTPUT)


if __name__ == "__main__":
    main()
